from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageOps
import os

# Função para criar uma tabela no documento
def adicionar_tabela(doc, detalhes):
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Total de Imagens'
    hdr_cells[1].text = 'Imagens com "alt"'
    hdr_cells[2].text = 'Imagens sem "alt"'

    # Preencher a tabela com os dados
    row_cells = table.add_row().cells
    row_cells[0].text = str(detalhes['total_imagens'])
    row_cells[1].text = str(detalhes['imagens_com_alt'])
    row_cells[2].text = str(detalhes['qtd_imagens_sem_alt'])

    doc.add_paragraph()

# Função para adicionar um quadro preto mais fino à imagem
def adicionar_quadro_preto(imagem_path):
    # Abrir a imagem
    img = Image.open(imagem_path)

    # Adicionar uma borda preta mais fina (3px de espessura)
    img_com_quadro = ImageOps.expand(img, border=3, fill='black')

    # Salvar a imagem com o quadro
    img_com_quadro_path = os.path.join("img", "quadro_" + os.path.basename(imagem_path))
    img_com_quadro.save(img_com_quadro_path)

    return img_com_quadro_path

# Função para adicionar um hyperlink clicável
def add_hyperlink(paragraph, url, text):
    """
    Adiciona um hyperlink clicável no parágrafo.
    :param paragraph: Parágrafo onde o hyperlink será inserido.
    :param url: URL que será vinculada.
    :param text: Texto que será clicável.
    """
    # Criar o relacionamento do hyperlink com o documento
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Criar o elemento <w:hyperlink> com o ID do relacionamento
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Criar o elemento de execução (run) para o texto do hyperlink
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')  # Formatação do run
    rStyle = OxmlElement('w:rStyle')  # Estilo de hyperlink
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text

    # Adicionar o novo run ao hyperlink
    hyperlink.append(new_run)

    # Adicionar o hyperlink ao parágrafo
    paragraph._p.append(hyperlink)

# Função para adicionar imagem e sua URL em duas colunas
def adicionar_imagens(doc, imagens):
    for imagem in imagens:
        nome = imagem['img_url'].split('/')[-1]  # Nome da imagem com base na URL

        table = doc.add_table(rows=1, cols=2)  # Criar uma tabela com 2 colunas

        # Adicionar a imagem à primeira coluna
        img_cell = table.rows[0].cells[0]
        p = img_cell.add_paragraph()
        run = p.add_run()
        # Se você não deseja incluir a imagem, remova esta linha:
        # run.add_picture(imagem_path, width=Inches(2))  # Adiciona a imagem com o quadro

        # Colocar a URL diretamente na segunda coluna
        url_cell = table.rows[0].cells[1]
        p_url = url_cell.add_paragraph()
        p_url.add_run(imagem['img_url'])  # Exibe a URL diretamente



# Função principal para gerar o relatório
def gerar_relatorio_auditoria(resultado_analises: dict, nome_arquivo='auditoria_relatorio.docx'):
    doc = Document()
    doc.add_heading('Relatório de Auditoria', 0)
    
    for url, detalhes in resultado_analises.items():
        doc.add_heading(f'URL Analisada: {url}', level=1)
        
        # Adicionar Tabela com Resumo de Imagens
        adicionar_tabela(doc, detalhes)
        
        # Adicionar imagens sem 'alt' com suas URLs em duas colunas
        adicionar_imagens(doc, detalhes['detalhes_imagens_sem_alt'])
        
    doc.save(nome_arquivo)
    print(f"Relatório de auditoria gerado: {nome_arquivo}")

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
# Função para adicionar bordas e cor de fundo à tabela
def adicionar_bordas(tabela):
    for row in tabela.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')

            # Definir as bordas
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')  # Espessura da borda
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # Cor preta para a borda
                tcBorders.append(border)

            # Definir a cor de fundo
            cell._element.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="ABC3DF"/>'.format(nsdecls('w')))
            )

            tcPr.append(tcBorders)

def adicionar_imagens_formatado(doc, imagens):
    for idx, imagem in enumerate(imagens):
        nome_arquivo = imagem['img_url'].split('/')[-1]  # Nome da imagem com base na URL
        caminho_imagem = os.path.join("img", nome_arquivo)  # Caminho para a imagem

        # Criar uma tabela com 2 colunas
        table = doc.add_table(rows=2, cols=2)  # Duas linhas: uma para o cabeçalho, outra para conteúdo

        # Adicionar bordas à tabela
        adicionar_bordas(table)

        # Cabeçalhos
        table.cell(0, 0).text = f"Imagem {idx + 1}"
        table.cell(0, 1).text = "Texto alternativo"
        
        # Configurar estilo dos cabeçalhos
        table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
        table.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(12)
        table.cell(0, 1).paragraphs[0].runs[0].font.size = Pt(12)

        # Segunda linha: adicionar a imagem à primeira coluna
        img_cell = table.cell(1, 0)
        p_img = img_cell.add_paragraph()
        run_img = p_img.add_run()
        p_img.add_run(f"\n{imagem['img_url']}\n").font.color.rgb = RGBColor(0, 0, 255)  # URL da imagem em azul

        # Verifica se a imagem foi baixada corretamente
        if os.path.exists(caminho_imagem):
            try:
                run_img.add_picture(caminho_imagem, width=Inches(2))  # Adiciona a imagem
            except Exception as e:
                p_img.add_run(f"Erro ao adicionar a imagem: {e}")
        else:
            p_img.add_run("Imagem não disponível")

        # Adicionar o nome da imagem sem o hiperlink
        p_img.add_run(f"\n{nome_arquivo}").font.color.rgb = RGBColor(128, 0, 128)  # Nome da imagem em roxo
            
        # Adicionar texto alternativo na segunda coluna
        alt_text_cell = table.cell(1, 1)
        p_alt_text = alt_text_cell.add_paragraph()
        p_alt_text.add_run(imagem.get('alt_text', 'Texto alternativo não disponível'))

        doc.add_paragraph()  # Adicionar um espaço entre as tabelas


# Função para gerar o relatório em docx (sem hiperlink)
def gerar_relatorio_docx(resultado_analises: dict, nome_arquivo='relatorio_auditoria.docx'):
    doc = Document()
    doc.add_heading('Relatório de Auditoria de Imagens', 0)

    for url, detalhes in resultado_analises.items():
        doc.add_heading(f'URL Analisada: {url}', level=1)
        adicionar_imagens_formatado(doc, detalhes['detalhes_imagens_sem_alt'])

    doc.save(nome_arquivo)
    print(f"Relatório DOCX gerado: {nome_arquivo}")


# # Função para gerar o relatório em PDF
# def gerar_relatorio_pdf(resultado_analises, nome_arquivo='relatorio_auditoria.pdf'):
#     pdf = FPDF()
#     pdf.set_auto_page_break(auto=True, margin=15)

#     # Definir margens: esquerda, topo, direita
#     pdf.set_margins(left=15, top=20, right=15)

#     pdf.add_page()

#     pdf.set_font('Arial', 'B', 16)
#     pdf.cell(0, 10, 'Relatório de Auditoria de Imagens', ln=True, align='C')
#     pdf.ln(10)

#     for url, detalhes in resultado_analises.items():
#         pdf.set_font('Arial', 'B', 12)
#         pdf.cell(0, 10, f"URL Analisada: {url}", ln=True)

#         for idx, imagem in enumerate(detalhes['detalhes_imagens_sem_alt']):
#             nome = imagem['img_url'].split('/')[-1]  # Nome da imagem
#             caminho_imagem = os.path.join("img", nome)

#             # Adicionar texto da imagem
#             pdf.set_font('Arial', 'B', 12)
#             pdf.cell(0, 10, f"Imagem {idx + 1}: {nome}", ln=True)

#             # Verificar se a imagem existe e ajustar o tamanho com a margem definida
#             if os.path.exists(caminho_imagem):
#                 try:
#                     pdf.image(caminho_imagem, w=70, x=pdf.get_x(), y=pdf.get_y())
#                     pdf.ln(30)  # Ajusta a altura do parágrafo após a imagem
#                 except RuntimeError:
#                     pdf.cell(0, 10, "Erro ao adicionar a imagem.", ln=True)
#             else:
#                 pdf.cell(0, 10, "Imagem não disponível", ln=True)

#             # Adicionar texto alternativo
#             pdf.set_font('Arial', '', 12)
#             texto_alt = imagem.get('alt_text', 'Texto alternativo não disponível')
#             pdf.multi_cell(0, 10, f"Texto alternativo: {texto_alt}")
#             pdf.ln(5)

#     pdf.output(nome_arquivo)
#     print(f"Relatório PDF gerado: {nome_arquivo}")

